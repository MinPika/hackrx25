import os
import io
import json
import re
import time
import tempfile
import hashlib
import asyncio
from typing import List, Dict, Optional, Tuple, Any
from concurrent.futures import ThreadPoolExecutor
import logging

# FastAPI and web components
from fastapi import FastAPI, HTTPException, Depends, status, BackgroundTasks
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uvicorn

# High-performance document processing
import fitz  # PyMuPDF for ultra-fast PDF processing
import httpx  # Async HTTP client
import docx
import email
from email.mime.text import MIMEText

# Fast embeddings and vector search
import numpy as np
import faiss
from sentence_transformers import SentenceTransformer
from rank_bm25 import BM25Okapi

# LLM
from groq import Groq

# Optimized configuration
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Constants
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "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")
HACKRX_AUTH_TOKEN = "95f763f2e367cc7e5f72304cb9e9b84229f97f2a5b2b08f14b5034e8328596ec"

# Global models - loaded once at startup
embedding_model = None
groq_client = None

class QueryRequest(BaseModel):
    documents: str
    questions: List[str]

class QueryResponse(BaseModel):
    answers: List[str]

# Security
security = HTTPBearer()

def verify_token(credentials: HTTPAuthorizationCredentials = Depends(security)):
    if credentials.credentials != HACKRX_AUTH_TOKEN:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid authentication credentials"
        )
    return credentials.credentials

class OptimizedDocumentProcessor:
    """Ultra-fast document processor optimized for speed"""
    
    def __init__(self):
        self.http_client = httpx.AsyncClient(timeout=30.0)
        self.executor = ThreadPoolExecutor(max_workers=4)
    
    async def download_document(self, url: str) -> bytes:
        """Download document with async HTTP"""
        try:
            logger.info(f"Downloading document from URL...")
            response = await self.http_client.get(url)
            response.raise_for_status()
            logger.info(f"Downloaded {len(response.content)} bytes")
            return response.content
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Download failed: {str(e)}")
    
    def detect_file_type(self, content: bytes) -> str:
        """Fast file type detection"""
        if content.startswith(b'%PDF'):
            return 'pdf'
        elif content.startswith(b'PK'):
            return 'docx'
        elif b'From:' in content[:1000] and b'Subject:' in content[:1000]:
            return 'email'
        return 'unknown'
    
    def process_pdf_fast(self, content: bytes) -> str:
        """Ultra-fast PDF processing using PyMuPDF"""
        try:
            # Open PDF from memory buffer - no disk I/O
            doc = fitz.open(stream=content, filetype="pdf")
            
            text_blocks = []
            total_pages = len(doc)
            logger.info(f"Processing {total_pages} pages")
            
            # Process pages in parallel batches
            def process_page_batch(start_page: int, end_page: int) -> List[str]:
                batch_text = []
                for page_num in range(start_page, min(end_page, total_pages)):
                    page = doc[page_num]
                    
                    # Fast text extraction - no OCR, pure text
                    text = page.get_text()
                    
                    if text.strip():
                        # Clean and structure text
                        cleaned_text = self.clean_text(text)
                        if cleaned_text:
                            batch_text.append(f"=== PAGE {page_num + 1} ===\n{cleaned_text}")
                
                return batch_text
            
            # Process in batches of 5 pages
            batch_size = 5
            futures = []
            
            for i in range(0, total_pages, batch_size):
                future = self.executor.submit(process_page_batch, i, i + batch_size)
                futures.append(future)
            
            # Collect results
            for future in futures:
                text_blocks.extend(future.result())
            
            doc.close()
            
            full_text = "\n\n".join(text_blocks)
            logger.info(f"Extracted {len(full_text)} characters from {total_pages} pages")
            return full_text
            
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"PDF processing failed: {str(e)}")
    
    def process_docx_fast(self, content: bytes) -> str:
        """Fast DOCX processing"""
        try:
            with tempfile.NamedTemporaryFile() as tmp_file:
                tmp_file.write(content)
                tmp_file.flush()
                
                doc = docx.Document(tmp_file.name)
                text_blocks = []
                
                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        text_blocks.append(text)
                
                # Process tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            text_blocks.append(row_text)
                
                return "\n\n".join(text_blocks)
                
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"DOCX processing failed: {str(e)}")
    
    def process_email_fast(self, content: bytes) -> str:
        """Fast email processing"""
        try:
            content_str = content.decode('utf-8', errors='ignore')
            msg = email.message_from_string(content_str)
            
            text_parts = []
            
            # Headers
            text_parts.append(f"From: {msg.get('From', 'Unknown')}")
            text_parts.append(f"To: {msg.get('To', 'Unknown')}")
            text_parts.append(f"Subject: {msg.get('Subject', 'No Subject')}")
            text_parts.append(f"Date: {msg.get('Date', 'Unknown')}")
            text_parts.append("="*50)
            
            # Body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        body = part.get_payload(decode=True)
                        if body:
                            text_parts.append(body.decode('utf-8', errors='ignore'))
            else:
                body = msg.get_payload(decode=True)
                if body:
                    text_parts.append(body.decode('utf-8', errors='ignore'))
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Email processing failed: {str(e)}")
    
    def clean_text(self, text: str) -> str:
        """Fast text cleaning"""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        # Remove page headers/footers patterns
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            # Skip common headers/footers
            if (len(line) < 3 or 
                line.isdigit() or 
                re.match(r'^Page \d+', line) or
                len(line) > 500):  # Skip very long lines (likely formatting issues)
                continue
            cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    async def process_document(self, url: str) -> str:
        """Main processing pipeline"""
        # Download
        content = await self.download_document(url)
        
        # Detect type
        file_type = self.detect_file_type(content)
        logger.info(f"Detected file type: {file_type}")
        
        # Process based on type
        if file_type == 'pdf':
            text = self.process_pdf_fast(content)
        elif file_type == 'docx':
            text = self.process_docx_fast(content)
        elif file_type == 'email':
            text = self.process_email_fast(content)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_type}")
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="No text extracted from document")
        
        return text

class UltraFastRetriever:
    """High-speed retrieval system using FAISS + BM25"""
    
    def __init__(self, embedding_model: SentenceTransformer):
        self.embedding_model = embedding_model
        self.chunks = []
        self.embeddings = None
        self.bm25 = None
        self.faiss_index = None
    
    def create_chunks(self, text: str) -> List[str]:
        """Smart chunking optimized for speed and accuracy"""
        # Split by double newlines first (paragraphs)
        paragraphs = text.split('\n\n')
        
        chunks = []
        current_chunk = ""
        target_size = 400  # tokens ≈ 300 words
        overlap_size = 50
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # If paragraph is small, accumulate
            if len(current_chunk) + len(para) < target_size:
                current_chunk += ("\n\n" if current_chunk else "") + para
            else:
                # Save current chunk if not empty
                if current_chunk:
                    chunks.append(current_chunk)
                    
                    # Start new chunk with overlap
                    overlap_text = current_chunk[-overlap_size:] if len(current_chunk) > overlap_size else ""
                    current_chunk = overlap_text + ("\n\n" if overlap_text else "") + para
                else:
                    current_chunk = para
        
        # Add final chunk
        if current_chunk:
            chunks.append(current_chunk)
        
        # Filter out very short chunks
        chunks = [chunk for chunk in chunks if len(chunk.strip()) > 50]
        
        logger.info(f"Created {len(chunks)} chunks")
        return chunks
    
    # def build_index(self, text: str):
    #     """Build retrieval index from text"""
    #     start_time = time.time()
        
    #     # Create chunks
    #     self.chunks = self.create_chunks(text)
        
    #     if not self.chunks:
    #         raise HTTPException(status_code=400, detail="No valid chunks created")
        
    #     # Create embeddings in batch
    #     logger.info("Creating embeddings...")
    #     self.embeddings = self.embedding_model.encode(
    #         self.chunks, 
    #         batch_size=32,
    #         show_progress_bar=False,
    #         convert_to_numpy=True
    #     )
        
    #     # Build FAISS index
    #     dimension = self.embeddings.shape[1]
    #     self.faiss_index = faiss.IndexFlatL2(dimension)
    #     self.faiss_index.add(self.embeddings.astype('float32'))
        
    #     # Build BM25 index
    #     tokenized_chunks = [chunk.lower().split() for chunk in self.chunks]
    #     self.bm25 = BM25Okapi(tokenized_chunks)
        
    #     build_time = time.time() - start_time
    #     logger.info(f"Built retrieval index in {build_time:.2f}s")
    def build_index(self, text: str):
        """Build retrieval index from text - optimized"""
        start_time = time.time()
        
        # Create chunks
        self.chunks = self.create_chunks(text)
        
        if not self.chunks:
            raise HTTPException(status_code=400, detail="No valid chunks created")
        
        # Create embeddings with optimizations
        logger.info("Creating embeddings...")
        
        # OPTIMIZATION 1: Larger batch size for faster processing
        # OPTIMIZATION 2: Normalize embeddings for faster similarity search
        # OPTIMIZATION 3: Use float32 to reduce memory and increase speed
        self.embeddings = self.embedding_model.encode(
            self.chunks, 
            batch_size=64,  # Increased from 32
            show_progress_bar=False,
            convert_to_numpy=True,
            normalize_embeddings=True,  # Faster cosine similarity
            device='cpu'
        ).astype(np.float32)  # Reduce memory usage
        
        # Build FAISS index with optimizations
        dimension = self.embeddings.shape[1]
        
        # Use IP (Inner Product) index since embeddings are normalized
        # This is faster than L2 distance for normalized vectors
        self.faiss_index = faiss.IndexFlatIP(dimension)
        self.faiss_index.add(self.embeddings)
        
        # Build BM25 index with optimizations
        # Pre-tokenize chunks for faster BM25 search
        tokenized_chunks = [chunk.lower().split() for chunk in self.chunks]
        self.bm25 = BM25Okapi(tokenized_chunks)
        
        build_time = time.time() - start_time
        logger.info(f"Built retrieval index in {build_time:.2f}s")
    
    # def hybrid_search(self, query: str, top_k: int = 5) -> List[Tuple[str, float]]:
    #     """Hybrid search using FAISS + BM25"""
    #     query_embedding = self.embedding_model.encode([query], convert_to_numpy=True)
        
    #     # FAISS semantic search
    #     distances, indices = self.faiss_index.search(query_embedding.astype('float32'), top_k * 2)
    #     semantic_results = [(self.chunks[idx], 1.0 / (1.0 + dist)) for idx, dist in zip(indices[0], distances[0])]
        
    #     # BM25 keyword search
    #     tokenized_query = query.lower().split()
    #     bm25_scores = self.bm25.get_scores(tokenized_query)
    #     bm25_results = [(self.chunks[i], score) for i, score in enumerate(bm25_scores)]
    #     bm25_results.sort(key=lambda x: x[1], reverse=True)
    #     bm25_results = bm25_results[:top_k * 2]
        
    #     # Combine using reciprocal rank fusion
    #     chunk_scores = {}
        
    #     # Add semantic scores
    #     for rank, (chunk, score) in enumerate(semantic_results):
    #         chunk_scores[chunk] = chunk_scores.get(chunk, 0) + 1.0 / (rank + 1)
        
    #     # Add BM25 scores
    #     for rank, (chunk, score) in enumerate(bm25_results):
    #         chunk_scores[chunk] = chunk_scores.get(chunk, 0) + 1.0 / (rank + 1)
        
    #     # Sort by combined score
    #     final_results = sorted(chunk_scores.items(), key=lambda x: x[1], reverse=True)
        
    #     return final_results[:top_k]
    def hybrid_search(self, query: str, top_k: int = 5) -> List[Tuple[str, float]]:
        """Hybrid search using FAISS + BM25 - optimized"""
        # Normalize query embedding for IP search
        query_embedding = self.embedding_model.encode(
            [query], 
            convert_to_numpy=True,
            normalize_embeddings=True,
            device='cpu'
        ).astype(np.float32)
        
        # FAISS semantic search (using Inner Product for normalized vectors)
        scores, indices = self.faiss_index.search(query_embedding, top_k * 2)
        semantic_results = [(self.chunks[idx], score) for idx, score in zip(indices[0], scores[0])]
        
        # BM25 keyword search
        tokenized_query = query.lower().split()
        bm25_scores = self.bm25.get_scores(tokenized_query)
        bm25_results = [(self.chunks[i], score) for i, score in enumerate(bm25_scores)]
        bm25_results.sort(key=lambda x: x[1], reverse=True)
        bm25_results = bm25_results[:top_k * 2]
        
        # Combine using reciprocal rank fusion
        chunk_scores = {}
        
        # Add semantic scores
        for rank, (chunk, score) in enumerate(semantic_results):
            chunk_scores[chunk] = chunk_scores.get(chunk, 0) + 1.0 / (rank + 1)
        
        # Add BM25 scores
        for rank, (chunk, score) in enumerate(bm25_results):
            chunk_scores[chunk] = chunk_scores.get(chunk, 0) + 1.0 / (rank + 1)
        
        # Sort by combined score
        final_results = sorted(chunk_scores.items(), key=lambda x: x[1], reverse=True)
        
        return final_results[:top_k]
class SmartQAGenerator:
    """Intelligent QA generation with domain-adaptive prompts"""
    
    def __init__(self, groq_client: Groq):
        self.groq_client = groq_client
    
    def classify_question_type(self, question: str) -> str:
        """Classify question to adapt prompt"""
        question_lower = question.lower()
        
        if any(word in question_lower for word in ['cover', 'coverage', 'covered', 'eligible', 'benefit']):
            return 'coverage'
        elif any(word in question_lower for word in ['period', 'waiting', 'grace', 'time', 'duration']):
            return 'time_period'
        elif any(word in question_lower for word in ['limit', 'amount', 'percentage', '%', 'discount', 'premium']):
            return 'financial'
        elif any(word in question_lower for word in ['define', 'definition', 'meaning', 'what is']):
            return 'definition'
        elif any(word in question_lower for word in ['condition', 'requirement', 'criteria', 'eligibility']):
            return 'conditions'
        else:
            return 'general'
    
    def create_adaptive_prompt(self, question: str, context: str, question_type: str) -> str:
        """Create prompts adapted to question type"""
        
        base_instructions = """You are an expert document analyst. Answer questions based STRICTLY on the provided context.

CRITICAL RULES:
1. Base your answer ONLY on the provided context
2. If information is not in the context, state "not specified in the document"
3. Be precise and cite specific policy terms when available
4. Keep answers concise but comprehensive"""
        
        type_specific_instructions = {
            'coverage': "Focus on what is/isn't covered, any exclusions, and specific conditions for coverage.",
            'time_period': "Identify specific time periods, waiting periods, grace periods, and any related conditions.",
            'financial': "Extract exact amounts, percentages, limits, and financial terms. Include any calculation methods.",
            'definition': "Provide the exact definition as stated in the document. Include any specific criteria or characteristics.",
            'conditions': "List all conditions, requirements, or criteria that must be met.",
            'general': "Provide a comprehensive answer addressing all aspects of the question."
        }
        
        prompt = f"""{base_instructions}

QUESTION TYPE: {question_type}
FOCUS: {type_specific_instructions.get(question_type, type_specific_instructions['general'])}

CONTEXT:
{context}

QUESTION: {question}

Provide a clear, accurate answer based on the context above:"""
        
        return prompt
    
    def generate_answer(self, question: str, context_chunks: List[Tuple[str, float]]) -> str:
        """Generate answer using Groq"""
        try:
            # Combine top chunks
            context = "\n\n---\n\n".join([chunk for chunk, score in context_chunks])
            
            # Classify question
            question_type = self.classify_question_type(question)
            
            # Create adaptive prompt
            prompt = self.create_adaptive_prompt(question, context, question_type)
            
            # Generate answer
            response = self.groq_client.chat.completions.create(
                model="meta-llama/llama-4-scout-17b-16e-instruct",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                max_tokens=500,
                stream=False
            )
            
            return response.choices[0].message.content.strip()
            
        except Exception as e:
            logger.error(f"Answer generation failed: {e}")
            return f"Unable to process the question due to technical issues: {str(e)}"

class OptimizedQASystem:
    """Main QA system orchestrating all components"""
    
    def __init__(self):
        self.doc_processor = OptimizedDocumentProcessor()
        self.retriever = None
        self.qa_generator = SmartQAGenerator(groq_client)
    
    async def process_queries(self, document_url: str, questions: List[str]) -> List[str]:
        """Main processing pipeline optimized for speed"""
        start_time = time.time()
        
        try:
            # Process document
            logger.info("Processing document...")
            text = await self.doc_processor.process_document(document_url)
            
            # Build retrieval index
            logger.info("Building retrieval index...")
            self.retriever = UltraFastRetriever(embedding_model)
            self.retriever.build_index(text)
            
            # Process all questions
            answers = []
            for i, question in enumerate(questions, 1):
                logger.info(f"Processing question {i}/{len(questions)}")
                
                # Retrieve relevant chunks
                relevant_chunks = self.retriever.hybrid_search(question, top_k=3)
                
                # Generate answer
                answer = self.qa_generator.generate_answer(question, relevant_chunks)
                answers.append(answer)
            
            total_time = time.time() - start_time
            logger.info(f"Processed {len(questions)} questions in {total_time:.2f}s")
            
            return answers
            
        except Exception as e:
            logger.error(f"Processing failed: {e}")
            raise HTTPException(status_code=500, detail=str(e))

# FastAPI application
app = FastAPI(
    title="HackRx 6.0 - Optimized Document QA System",
    description="Ultra-fast document processing and question answering",
    version="2.0.0"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Global QA system
qa_system = None

@app.on_event("startup")
async def startup_event():
    """Initialize models on startup"""
    global qa_system, embedding_model, groq_client
    
    try:
        logger.info("Starting optimized HackRx system...")
        
        # Initialize embedding model
        logger.info("Loading embedding model...")
        embedding_model = SentenceTransformer('BAAI/bge-base-en-v1.5', device='cpu')
        
        # Initialize Groq client
        logger.info("Initializing Groq client...")
        groq_client = Groq(api_key=GROQ_API_KEY)
        
        # Initialize QA system
        qa_system = OptimizedQASystem()
        
        logger.info("✅ System ready!")
        
    except Exception as e:
        logger.error(f"Startup failed: {e}")
        raise

@app.get("/")
async def root():
    """System status"""
    return {
        "status": "online",
        "system": "HackRx 6.0 - Optimized Document QA",
        "version": "2.0.0",
        "optimizations": [
            "PyMuPDF for ultra-fast PDF processing",
            "FAISS + BM25 hybrid search",
            "BGE embeddings for accuracy",
            "Groq for lightning-fast inference",
            "Adaptive prompting system",
            "Parallel processing pipeline"
        ]
    }

@app.get("/health")
async def health_check():
    """Detailed health check"""
    return {
        "status": "healthy",
        "timestamp": time.time(),
        "components": {
            "qa_system": "✅" if qa_system else "❌",
            "embedding_model": "✅" if embedding_model else "❌",
            "groq_client": "✅" if groq_client else "❌",
        }
    }

@app.post("/hackrx/run", response_model=QueryResponse)
async def hackrx_run(
    request: QueryRequest,
    token: str = Depends(verify_token)
) -> QueryResponse:
    """Optimized HackRx endpoint"""
    try:
        logger.info(f"Received request with {len(request.questions)} questions")
        
        if qa_system is None:
            raise HTTPException(status_code=500, detail="QA system not initialized")
        
        # Process queries
        answers = await qa_system.process_queries(request.documents, request.questions)
        
        logger.info("Successfully processed all questions")
        return QueryResponse(answers=answers)
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error: {e}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")

if __name__ == "__main__":
    print("🚀 HackRx 6.0 - Ultra-Fast Optimized System")
    print("=" * 50)
    print("⚡ Optimizations:")
    print("  • PyMuPDF: 10x faster PDF processing")
    print("  • No OCR: Pure text extraction")
    print("  • FAISS: Microsecond vector search")
    print("  • BM25: Keyword matching")
    print("  • BGE: Accurate embeddings")
    print("  • Groq: Lightning inference")
    print("  • Parallel processing")
    print("=" * 50)
    
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=8000,
        reload=False,
        log_level="info"
    )